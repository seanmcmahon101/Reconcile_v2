# Flask application for Excel file analysis, formula creation, and account reconciliation
# Using Google's Gemini AI for intelligent processing and analysis
# Add tracking for user activities with a simple in-memory store
# This could be expanded to use a database in production

# At the top of app.py with other imports
from collections import deque
import time

# Add activity tracking
user_activities = {}
MAX_ACTIVITIES = 10

def track_activity(user_id, activity_type, details=None):
    """Tracks user activity for the dashboard."""
    if user_id not in user_activities:
        user_activities[user_id] = deque(maxlen=MAX_ACTIVITIES)

    user_activities[user_id].appendleft({
        'type': activity_type,
        'timestamp': time.time(),
        'details': details or {}
    })

# Then modify your routes to track activities
@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    """Handles the main application logic or redirects to dashboard."""
    # For a logged-in user without active analysis, redirect to dashboard
    if request.method == 'GET' and not session.get('current_explanation_html'):
        return redirect(url_for('dashboard'))

    # Rest of your existing index route code here
    # Add tracking when analysis is completed:
    if explanation_html:
        track_activity(
            current_user.id,
            'analysis',
            {'filename': secure_filename(file.filename)}
        )
    # ...

# Add a new dashboard route
@app.route('/dashboard')
@login_required
def dashboard():
    """Displays the user dashboard with recent activities and system status."""
    # Get user's recent activities
    activities = []
    if current_user.id in user_activities:
        raw_activities = list(user_activities[current_user.id])

        for activity in raw_activities:
            activity_type = activity['type']
            details = activity['details']
            timestamp = datetime.datetime.fromtimestamp(activity['timestamp'])

            if activity_type == 'analysis':
                activities.append({
                    'title': 'Sheet Analysis',
                    'description': f"Analyzed {details.get('filename', 'a sheet')}",
                    'time': timestamp.strftime('%Y-%m-%d %H:%M'),
                    'icon': 'fa-file-excel',
                    'icon_class': 'bg-light-primary text-primary'
                })
            elif activity_type == 'formula':
                activities.append({
                    'title': 'Formula Creation',
                    'description': details.get('description', 'Created a formula'),
                    'time': timestamp.strftime('%Y-%m-%d %H:%M'),
                    'icon': 'fa-calculator',
                    'icon_class': 'bg-light-success text-success'
                })
            elif activity_type == 'reconciliation':
                activities.append({
                    'title': 'Account Reconciliation',
                    'description': f"Reconciled {details.get('file1', 'Sheet 1')} with {details.get('file2', 'Sheet 2')}",
                    'time': timestamp.strftime('%Y-%m-%d %H:%M'),
                    'icon': 'fa-balance-scale',
                    'icon_class': 'bg-light-warning text-warning'
                })

    # Get system stats
    api_status, _ = test_api_connection()

    # Placeholder stats (in production, you'd get these from a database)
    stats = {
        'analyses_count': sum(1 for acts in user_activities.get(current_user.id, [])
                            if acts['type'] == 'analysis'),
        'formulas_count': sum(1 for acts in user_activities.get(current_user.id, [])
                             if acts['type'] == 'formula'),
        'reconciliations_count': sum(1 for acts in user_activities.get(current_user.id, [])
                                   if acts['type'] == 'reconciliation')
    }

    return render_template(
        'dashboard.html',
        stats=stats,
        recent_activities=activities,
        system_status={
            'api_connected': api_status,
            'current_model': DEFAULT_MODEL_NAME,
            'last_update': datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
        },
        current_user=current_user
    )

import os
import logging
import datetime
from flask import Flask, render_template, request, redirect, url_for, session, send_file, flash, make_response, jsonify
from werkzeug.utils import secure_filename
import openpyxl
import google.generativeai as genai
import markdown
from docx import Document
from io import BytesIO

# Import for password protection
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from functools import wraps

app = Flask(__name__)
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'your_default_secret_key')
app.config['SESSION_COOKIE_SECURE'] = True  # Use secure cookies in production
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Prevent JavaScript access to session cookie
app.config['PERMANENT_SESSION_LIFETIME'] = datetime.timedelta(hours=2)  # Session timeout

# --- Configuration ---
API_KEY = os.getenv("GEMINI_API_KEY")
DEFAULT_MODEL_NAME = 'gemini-2.0-flash'
THINKING_MODEL_NAME = 'gemini-2.0-flash-thinking-exp'

# Enhanced system prompts for better LLM responses
SYSTEM_PROMPT = """You are ExcelAnalyst, an expert in analyzing Excel spreadsheets with deep knowledge of finance, accounting, and data analysis.

Your task is to provide a detailed, professional explanation of Excel files, addressing:

1. PURPOSE: What is the likely business purpose of this spreadsheet? (financial report, sales tracking, inventory management, etc.)
2. STRUCTURE: How is the data organized? Identify tables, sections, and their relationships.
3. DATA ANALYSIS: What key metrics, calculations, or trends are shown?
4. FORMULAS: Explain complex formulas in plain language, detailing their business logic and purpose.
5. KEY INSIGHTS: What are the most important takeaways a business user should understand?

Format your explanation with clear Markdown:
- Use ### for main section headings
- Use bullet points for lists of related items
- Use code blocks for formula examples
- Bold important terms or values
- Use tables when comparing multiple values

Assume your reader is a business professional who needs to quickly understand the spreadsheet's purpose and value.
"""

FORMULA_SYSTEM_PROMPT = """You are FormulaExpert, specializing in creating efficient, robust Excel formulas for financial and business applications.

When I describe what I need, provide:

1. OPTIMAL FORMULA: The most efficient, production-ready Excel formula for my needs
2. EXPLANATION: A clear explanation of how the formula works, component by component
3. EXAMPLE: A practical example with sample data showing input and expected output
4. ALTERNATIVES: If relevant, mention 1-2 alternative approaches with their pros/cons
5. BEST PRACTICES: Any tips for formula maintenance, error handling, or performance

For complex formulas, break them down into logical parts before explaining the whole solution.
Use proper Excel syntax and consider potential edge cases (errors, missing data, etc.).

Format your response with clear Markdown:
- Use ### headings for different sections
- Use code blocks for all formula examples
- Bold key concepts or functions
- Use a table if comparing approaches

Assume I am an intermediate Excel user familiar with basic functions but may need help with advanced techniques.
"""

RECONCILIATION_SYSTEM_PROMPT = """You are ReconciliationPro, a specialist in financial reconciliation and accounting using Excel data.

Your task is to compare two datasets and provide a detailed reconciliation analysis that would meet professional accounting standards. Focus on:

1. SUMMARY OF FINDINGS: Overall reconciliation status - balanced, partially reconciled, or significant discrepancies
2. KEY DISCREPANCIES: Identify specific differences between corresponding records, with exact values
3. PATTERN ANALYSIS: Identify any systematic issues (consistent time lags, rounding differences, missing categories)
4. ROOT CAUSES: Suggest likely reasons for discrepancies (timing differences, calculation errors, missing transactions)
5. RECOMMENDED ACTIONS: Provide specific next steps for full reconciliation
6. MATERIALITY ASSESSMENT: Evaluate if discrepancies exceed materiality thresholds for financial reporting

When analyzing:
- For financial data, calculate both absolute ($) and percentage (%) variances
- Flag entries present in only one dataset
- Identify matches that reconcile perfectly
- Note any structural or formatting differences that may cause mismatches

Format your analysis professionally with clear Markdown:
- Use ### for main section headings
- Bold critical discrepancies or large variances
- Use tables for reconciliation summaries and key numbers
- Use bullet points for recommendations
- Include totals and subtotals where appropriate

Your analysis must be precise, actionable, and reflect professional accounting standards.
"""

ADMIN_SYSTEM_PROMPT = """You are SystemAdmin, an expert in diagnosing and troubleshooting API integrations and system configurations.

Please analyze the provided system information and report on:

1. API CONNECTION STATUS: Verify if the API connection is working properly
2. CONFIGURATION ISSUES: Identify any missing or misconfigured settings
3. PERFORMANCE METRICS: Assess system response times and resource utilization
4. SECURITY CONCERNS: Flag any potential security vulnerabilities
5. RECOMMENDATIONS: Suggest specific improvements to system performance and reliability

Format your analysis with clear Markdown:
- Use ### for main section headings
- Use ✅ for successful checks and ❌ for failed checks
- Bold critical issues requiring immediate attention
- Use code blocks for configuration examples or commands
- Use tables for comparative metrics

Your report should be actionable and prioritize critical issues first.
"""

PROMPT_PREFIX = "The Excel sheet contains the following information in a structured way:\n"
PROMPT_SUFFIX = "\nAnalyze this Excel data comprehensively. Identify the business purpose, explain the data structure, interpret complex formulas, highlight key metrics, and provide insights a business user would find valuable. Use professional financial/accounting terminology where appropriate."

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'xlsx', 'xls'}
DEFAULT_DOCX_FILENAME = "excel_explanation.docx"
FORMULA_DOCX_FILENAME = "excel_formula.docx"
CHAT_DOCX_FILENAME = "excel_chat.docx"
RECONCILIATION_DOCX_FILENAME = "excel_reconciliation.docx"


os.makedirs(UPLOAD_FOLDER, exist_ok=True)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Flask-Login Configuration ---
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = "Please log in to access this page."

# --- User Data from Environment Variables ---
users = {}
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "admin_password")  # Default should be changed in production

# Add admin user
users[0] = {
    'username': ADMIN_USERNAME,
    'password_hash': generate_password_hash(ADMIN_PASSWORD),
    'is_admin': True
}

# Regular users
for i in range(1, 5):
    username = os.getenv(f"USER{i}_USERNAME")
    password = os.getenv(f"USER{i}_PASSWORD")
    if username and password:
        users[i] = {
            'username': username,
            'password_hash': generate_password_hash(password),
            'is_admin': False
        }
    else:
        logging.warning(f"User {i} credentials not fully configured via environment variables.")


class User(UserMixin):
    def __init__(self, id, username, password_hash, is_admin=False):
        self.id = id
        self.username = username
        self.password_hash = password_hash
        self.is_admin = is_admin

    def verify_password(self, password):
        return check_password_hash(self.password_hash, password)


@login_manager.user_loader
def load_user(user_id):
    user_data = users.get(int(user_id))
    if user_data:
        return User(
            id=user_id,
            username=user_data['username'],
            password_hash=user_data['password_hash'],
            is_admin=user_data.get('is_admin', False)
        )
    return None


def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            flash('You do not have permission to access this page.', 'error')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function


def configure_api():
    """Configures the Gemini API with the API key."""
    if not API_KEY:
        logging.error("API_KEY environment variable not set.")
        return False
    try:
        genai.configure(api_key=API_KEY)
        return True
    except Exception as e:
        logging.error(f"Error configuring Gemini API: {e}")
        return False


def test_api_connection():
    """Tests connection to Gemini API and returns status."""
    try:
        model = genai.GenerativeModel(DEFAULT_MODEL_NAME)
        response = model.generate_content("Hello, this is a test.")
        if response and response.text:
            return True, "API connection successful"
        return False, "API returned empty response"
    except Exception as e:
        return False, f"API connection failed: {str(e)}"


def allowed_file(filename):
    """Check if the uploaded file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def load_excel_data(file_path):
    """Loads data from the Excel file using openpyxl."""
    try:
        wb = openpyxl.load_workbook(file_path, data_only=False)
        sheet = wb.active
        return sheet
    except Exception as e:
        logging.error(f"Error loading Excel file: {e}")
        return None


def build_prompt_reconciliation(sheet1, sheet2):
    """Builds the prompt for Gemini API for reconciliation, comparing two sheets."""
    prompt_content = "Sheet 1 Data:\n"
    
    # Get header row from first sheet for column names
    headers1 = []
    for cell in sheet1[1]:
        headers1.append(str(cell.value) if cell.value is not None else "Unnamed Column")
    
    prompt_content += "Headers: " + ", ".join(headers1) + "\n\n"
    
    # Get data rows from first sheet
    for row_idx, row in enumerate(sheet1.iter_rows(min_row=2, max_row=sheet1.max_row, min_col=1, max_col=sheet1.max_column), start=2):
        row_values = []
        for cell in row:
            # Preserve formula information
            if cell.data_type == 'f':
                cell_value = f"FORMULA: {cell.value} = {cell.value}"
            else:
                cell_value = str(cell.value) if cell.value is not None else "None"
            row_values.append(cell_value)
        prompt_content += f"- Row {row_idx}: " + ", ".join(row_values) + "\n"

    # Similar processing for second sheet
    prompt_content += "\nSheet 2 Data:\n"
    
    # Get header row from second sheet
    headers2 = []
    for cell in sheet2[1]:
        headers2.append(str(cell.value) if cell.value is not None else "Unnamed Column")
    
    prompt_content += "Headers: " + ", ".join(headers2) + "\n\n"
    
    # Get data rows from second sheet
    for row_idx, row in enumerate(sheet2.iter_rows(min_row=2, max_row=sheet2.max_row, min_col=1, max_col=sheet2.max_column), start=2):
        row_values = []
        for cell in row:
            if cell.data_type == 'f':
                cell_value = f"FORMULA: {cell.value} = {cell.value}"
            else:
                cell_value = str(cell.value) if cell.value is not None else "None"
            row_values.append(cell_value)
        prompt_content += f"- Row {row_idx}: " + ", ".join(row_values) + "\n"

    # Add metadata about the sheets to help with reconciliation
    prompt_content += "\nReconciliation Metadata:\n"
    prompt_content += f"- Sheet 1 has {sheet1.max_row} rows and {sheet1.max_column} columns\n"
    prompt_content += f"- Sheet 2 has {sheet2.max_row} rows and {sheet2.max_column} columns\n"
    
    # Identify common column names for potential matching keys
    common_headers = set(headers1).intersection(set(headers2))
    prompt_content += f"- Common column headers: {', '.join(common_headers)}\n"
    
    # Check for potential ID columns
    id_columns = [header for header in common_headers if 'id' in header.lower() or 'code' in header.lower() or 'key' in header.lower()]
    if id_columns:
        prompt_content += f"- Potential matching key columns: {', '.join(id_columns)}\n"
    
    # Check for date columns
    date_columns = [header for header in common_headers if 'date' in header.lower() or 'time' in header.lower()]
    if date_columns:
        prompt_content += f"- Date-related columns: {', '.join(date_columns)}\n"
    
    # Check for amount columns
    amount_columns = [header for header in common_headers if 'amount' in header.lower() or 'value' in header.lower() or 'total' in header.lower()]
    if amount_columns:
        prompt_content += f"- Financial amount columns: {', '.join(amount_columns)}\n"

    full_prompt = RECONCILIATION_SYSTEM_PROMPT + "\n\nData from Sheet 1 and Sheet 2 to reconcile:\n" + prompt_content
    logging.info("Reconciliation prompt built successfully.")
    return full_prompt


def build_prompt(sheet):
    """Builds the prompt for the Gemini API based on the Excel sheet data."""
    prompt_content = ""
    
    # Extract sheet metadata
    prompt_content += f"Excel Sheet Metadata:\n"
    prompt_content += f"- Sheet name: {sheet.title}\n"
    prompt_content += f"- Total rows: {sheet.max_row}\n"
    prompt_content += f"- Total columns: {sheet.max_column}\n\n"
    
    # Get column headers (assuming first row contains headers)
    headers = []
    for cell in sheet[1]:
        header_value = str(cell.value) if cell.value is not None else "Unnamed Column"
        headers.append(header_value)
        
    prompt_content += f"Column Headers: {', '.join(headers)}\n\n"
    
    # Sample data - first 5 rows for context
    prompt_content += "Sample Data (First 5 rows):\n"
    for row_idx, row in enumerate(sheet.iter_rows(min_row=2, max_row=min(sheet.max_row, 6), min_col=1, max_col=sheet.max_column), start=2):
        row_values = []
        for cell in row:
            if cell.value is not None:
                row_values.append(f"{headers[cell.column-1]}: {cell.value}")
        if row_values:
            prompt_content += f"- Row {row_idx}: {', '.join(row_values)}\n"
    
    prompt_content += "\nDetailed Cell Information:\n"
    
    # Track formulas for special attention
    formulas = []
    
    # Process all non-empty cells
    for row in sheet.iter_rows(min_row=1, max_row=sheet.max_row, min_col=1, max_col=sheet.max_column):
        for cell in row:
            if cell.value is not None or cell.comment is not None:
                cell_info = ""
                
                # Handle formulas with special attention
                if cell.data_type == 'f':
                    formulas.append((cell.coordinate, cell.value))
                    cell_info = f"formula '{cell.value}'"
                elif cell.value is not None:
                    cell_info = f"value '{cell.value}'"
                else:
                    cell_info = "no value"

                # Include cell comments
                comment_text = ""
                if cell.comment:
                    comment_text_raw = cell.comment.text.strip()
                    comment_text_processed = comment_text_raw.replace('\n', ' ')
                    comment_text = f" with comment '{comment_text_processed}'"

                prompt_content += f"- Cell {cell.coordinate} has {cell_info}{comment_text}.\n"
    
    # Add a special section for formulas if any were found
    if formulas:
        prompt_content += "\nFormula Analysis:\n"
        for coord, formula in formulas:
            prompt_content += f"- Formula at {coord}: {formula}\n"
    
    # Check for specific financial patterns
    financial_indicators = []
    for header in headers:
        lower_header = header.lower()
        if any(term in lower_header for term in ['total', 'sum', 'balance', 'account', 'revenue', 'expense', 'profit', 'loss', 'asset', 'liability']):
            financial_indicators.append(header)
    
    if financial_indicators:
        prompt_content += "\nPotential Financial Indicators Found:\n"
        prompt_content += f"- This appears to contain financial data with these indicators: {', '.join(financial_indicators)}\n"

    full_prompt = PROMPT_PREFIX + prompt_content + PROMPT_SUFFIX
    logging.info("Prompt built successfully.")
    return full_prompt


def get_explanation_from_gemini(prompt, model_name):
    """Gets explanation from Gemini API."""
    model = genai.GenerativeModel(model_name)
    try:
        response = model.generate_content(
            prompt, 
            generation_config=genai.types.GenerationConfig(
                temperature=0.2,
                top_p=0.95,
                top_k=40
            )
        )
        explanation = response.text
        logging.info(f"Explanation received from Gemini API using model: {model_name}")
        return explanation
    except Exception as e:
        logging.error(f"Error communicating with Gemini API: {e}")
        return None


def get_formula_from_gemini(prompt):
    """Gets formula from Gemini API using formula system prompt."""
    model = genai.GenerativeModel(DEFAULT_MODEL_NAME)
    full_prompt = FORMULA_SYSTEM_PROMPT + "\n\n" + prompt
    try:
        response = model.generate_content(
            full_prompt, 
            generation_config=genai.types.GenerationConfig(
                temperature=0.4,
                top_p=0.95,
                top_k=40
            )
        )
        formula_explanation = response.text
        logging.info("Formula explanation received from Gemini API.")
        return formula_explanation
    except Exception as e:
        logging.error(f"Error communicating with Gemini API for formula: {e}")
        return None


def export_to_docx(explanation, filename=DEFAULT_DOCX_FILENAME):
    """Exports content to a DOCX file in memory and returns BytesIO object."""
    doc = Document()
    
    # Add a title
    doc.add_heading(filename.replace('_', ' ').replace('.docx', '').title(), 0)
    
    # Add date
    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    # Add a horizontal line
    doc.add_paragraph().add_run().add_break()
    
    # Add the content with better formatting
    for line in explanation.splitlines():
        # Check if line is a heading (starts with #)
        if line.strip().startswith('# '):
            doc.add_heading(line.strip().replace('# ', ''), 1)
        elif line.strip().startswith('## '):
            doc.add_heading(line.strip().replace('## ', ''), 2)
        elif line.strip().startswith('### '):
            doc.add_heading(line.strip().replace('### ', ''), 3)
        elif line.strip().startswith('#### '):
            doc.add_heading(line.strip().replace('#### ', ''), 4)
        # Check if line is a bullet point
        elif line.strip().startswith('- '):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        # Check if line is a numbered list
        elif line.strip().startswith('1. ') or line.strip().startswith('2. '):
            doc.add_paragraph(line.strip()[3:], style='List Number')
        # Regular text
        else:
            doc.add_paragraph(line)

    docx_stream = BytesIO()
    try:
        doc.save(docx_stream)
        docx_stream.seek(0)
        logging.info(f"Content exported to DOCX in memory as {filename}.")
        return docx_stream
    except Exception as e:
        logging.error(f"Error exporting to DOCX: {e}")
        return None


# --- Routes ---
@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page."""
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        user_data = None
        user_id_found = None
        
        # Search for matching username in users dictionary
        for user_id, data in users.items():
            if data['username'] == username:
                user_data = data
                user_id_found = user_id
                break

        if user_data and check_password_hash(user_data['password_hash'], password):
            user = User(
                id=user_id_found, 
                username=username, 
                password_hash=user_data['password_hash'],
                is_admin=user_data.get('is_admin', False)
            )
            login_user(user)
            flash('Logged in successfully.')
            next_page = request.args.get('next')
            return redirect(next_page or url_for('index'))
        else:
            flash('Invalid username or password', 'error')
    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    """Logs out the current user."""
    logout_user()
    flash('Logged out successfully.')
    return redirect(url_for('index'))


@app.route('/', methods=['GET', 'POST'])
@login_required
def index():
    """Handles the main application logic for Excel sheet explanation."""
    explanation_html = None
    error = None
    model_name = DEFAULT_MODEL_NAME

    if request.method == 'POST':
        if 'excel_file' not in request.files:
            error = 'No file part'
        elif request.files['excel_file'].filename == '':
            error = 'No selected file'
        elif 'excel_file' in request.files and allowed_file(request.files['excel_file'].filename):
            file = request.files['excel_file']
            try:
                filename = secure_filename(file.filename)
                file_path = os.path.join(UPLOAD_FOLDER, filename)
                file.save(file_path)

                selected_model = request.form.get('model_select')
                if selected_model == 'thinking':
                    model_name = THINKING_MODEL_NAME
                else:
                    model_name = DEFAULT_MODEL_NAME

                sheet = load_excel_data(file_path)
                if sheet:
                    prompt = build_prompt(sheet)
                    explanation_markdown = get_explanation_from_gemini(prompt, model_name)

                    if explanation_markdown:
                        explanation_html = markdown.markdown(explanation_markdown)
                        session['explanation_markdown'] = explanation_markdown
                        session['current_explanation_html'] = explanation_html
                    else:
                        error = "Failed to get explanation from Gemini API."
                else:
                    error = "Failed to load Excel data."
            except Exception as e:
                error = f"An error occurred: {e}"
                logging.error(f"Error in index route: {e}")
            finally:
                if os.path.exists(file_path):
                    os.remove(file_path)
        else:
            error = 'Invalid file type. Allowed types are xlsx, xls'

    response = make_response(render_template(
        'index.html', 
        explanation_html=explanation_html, 
        error=error, 
        model_name=model_name, 
        current_user=current_user
    ))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response


@app.route('/export_docx')
@login_required
def export_docx_route():
    """Exports the explanation to DOCX format and allows download."""
    explanation_markdown = session.get('explanation_markdown')
    if not explanation_markdown:
        flash('No explanation available to export.', 'error')
        return redirect(url_for('index'))

    docx_stream = export_to_docx(explanation_markdown, DEFAULT_DOCX_FILENAME)
    if docx_stream:
        return send_file(
            docx_stream,
            as_attachment=True,
            download_name=DEFAULT_DOCX_FILENAME,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        flash('Error exporting to DOCX.', 'error')
        return redirect(url_for('index'))


@app.route('/formula_creator', methods=['GET', 'POST'])
@login_required
def formula_creator():
    """Handles the formula creation page."""
    formula_explanation_html = None
    error = None

    if request.method == 'POST':
        formula_description = request.form.get('formula_description')
        if formula_description:
            formula_explanation_markdown = get_formula_from_gemini(formula_description)
            if formula_explanation_markdown:
                formula_explanation_html = markdown.markdown(formula_explanation_markdown)
                session['formula_explanation_markdown'] = formula_explanation_markdown
            else:
                error = "Failed to get formula explanation from Gemini API."
        else:
            error = "Please enter a description for the formula you need."

    response = make_response(render_template(
        'formula_creator.html', 
        formula_explanation_html=formula_explanation_html, 
        error=error, 
        current_user=current_user
    ))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response


@app.route('/export_formula_docx')
@login_required
def export_formula_docx_route():
    """Exports the formula explanation to DOCX format."""
    formula_explanation_markdown = session.get('formula_explanation_markdown')
    if not formula_explanation_markdown:
        flash('No formula explanation available to export.', 'error')
        return redirect(url_for('formula_creator'))

    docx_stream = export_to_docx(formula_explanation_markdown, FORMULA_DOCX_FILENAME)
    if docx_stream:
        return send_file(
            docx_stream,
            as_attachment=True,
            download_name=FORMULA_DOCX_FILENAME,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        flash('Error exporting formula explanation to DOCX.', 'error')
        return redirect(url_for('formula_creator'))


@app.route('/chat', methods=['GET', 'POST'])
@login_required
def chat():
    """Handles the chat functionality after sheet analysis."""
    explanation_html = session.get('current_explanation_html')
    chat_history = session.get('chat_history', [])
    error = None

    if explanation_html is None:
        flash('Please analyze an Excel sheet first.', 'info')
        return redirect(url_for('index'))

    if request.method == 'POST':
        user_message = request.form.get('chat_message')
        if user_message:
            # Create a context-aware prompt for the chat
            prompt_context = f"""The user has analyzed an Excel sheet with the following explanation:
            
{session.get('explanation_markdown')}

CHAT HISTORY:
{' '.join([f"User: {m['user']} Bot: {m['bot']}" for m in chat_history])}

User's new question: {user_message}

Please provide a helpful, specific response focused on answering the user's question about the Excel sheet. 
Use Markdown formatting for clarity. If the question can't be answered based on the provided information, 
politely explain what additional details would be needed."""

            llm_response_markdown = get_explanation_from_gemini(prompt_context, DEFAULT_MODEL_NAME)
            if llm_response_markdown:
                llm_response_html = markdown.markdown(llm_response_markdown)
                chat_history.append({'user': user_message, 'bot': llm_response_html})
                session['chat_history'] = chat_history
            else:
                error = "Failed to get chat response from Gemini API."
        else:
            error = "Please enter a chat message."

    response = make_response(render_template(
        'chat.html', 
        explanation_html=explanation_html, 
        chat_history=chat_history, 
        error=error, 
        current_user=current_user
    ))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response


@app.route('/export_chat_docx')
@login_required
def export_chat_docx_route():
    """Exports the chat history to DOCX format."""
    chat_history = session.get('chat_history')
    if not chat_history:
        flash('No chat history available to export.', 'error')
        return redirect(url_for('chat'))

    chat_markdown = "# Excel Analysis Chat History\n\n"
    
    # Add the original explanation for context
    if session.get('explanation_markdown'):
        chat_markdown += "## Original Excel Analysis\n\n"
        chat_markdown += session.get('explanation_markdown') + "\n\n"
    
    chat_markdown += "## Chat History\n\n"
    for i, message in enumerate(chat_history, 1):
        chat_markdown += f"### Conversation {i}\n\n"
        chat_markdown += f"**User Question:**\n\n{message['user']}\n\n"
        chat_markdown += f"**Analysis Response:**\n\n{message['bot']}\n\n"
        chat_markdown += "---\n\n"

    docx_stream = export_to_docx(chat_markdown, CHAT_DOCX_FILENAME)
    if docx_stream:
        return send_file(
            docx_stream,
            as_attachment=True,
            download_name=CHAT_DOCX_FILENAME,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        flash('Error exporting chat history to DOCX.', 'error')
        return redirect(url_for('chat'))


@app.route('/reconcile', methods=['GET', 'POST'])
@login_required
def reconcile():
    """Handles the accounts reconciliation page and logic."""
    reconciliation_explanation_html = None
    error = None
    model_name = DEFAULT_MODEL_NAME

    if request.method == 'POST':
        if 'excel_file_1' not in request.files or 'excel_file_2' not in request.files:
            error = 'Need to upload both Sheet 1 and Sheet 2'
        elif request.files['excel_file_1'].filename == '' or request.files['excel_file_2'].filename == '':
            error = 'Both Sheet 1 and Sheet 2 files need to be selected'
        elif 'excel_file_1' in request.files and allowed_file(request.files['excel_file_1'].filename) and 'excel_file_2' in request.files and allowed_file(request.files['excel_file_2'].filename):
            file1 = request.files['excel_file_1']
            file2 = request.files['excel_file_2']
            file_path_1 = os.path.join(UPLOAD_FOLDER, secure_filename(file1.filename))
            file_path_2 = os.path.join(UPLOAD_FOLDER, secure_filename(file2.filename))

            try:
                file1.save(file_path_1)
                file2.save(file_path_2)

                selected_model = request.form.get('model_select')
                if selected_model == 'thinking':
                    model_name = THINKING_MODEL_NAME
                else:
                    model_name = DEFAULT_MODEL_NAME

                sheet1 = load_excel_data(file_path_1)
                sheet2 = load_excel_data(file_path_2)

                if sheet1 and sheet2:
                    prompt = build_prompt_reconciliation(sheet1, sheet2)
                    reconciliation_markdown = get_explanation_from_gemini(prompt, model_name)

                    if reconciliation_markdown:
                        reconciliation_explanation_html = markdown.markdown(reconciliation_markdown)
                        session['reconciliation_explanation_markdown'] = reconciliation_markdown
                    else:
                        error = "Failed to get reconciliation explanation from Gemini API."
                else:
                    error = "Failed to load data from one or both Excel files."
            except Exception as e:
                error = f"An error occurred during reconciliation: {e}"
                logging.error(f"Error in reconcile route: {e}")
            finally:
                if os.path.exists(file_path_1):
                    os.remove(file_path_1)
                if os.path.exists(file_path_2):
                    os.remove(file_path_2)
        else:
            error = 'Invalid file types. Allowed types are xlsx, xls for both sheets.'

    response = make_response(render_template(
        'reconcile.html', 
        reconciliation_explanation_html=reconciliation_explanation_html, 
        error=error, 
        current_user=current_user, 
        model_name=model_name
    ))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response


@app.route('/export_reconciliation_docx')
@login_required
def export_reconciliation_docx_route():
    """Exports the reconciliation explanation to DOCX format."""
    reconciliation_explanation_markdown = session.get('reconciliation_explanation_markdown')
    if not reconciliation_explanation_markdown:
        flash('No reconciliation explanation available to export.', 'error')
        return redirect(url_for('reconcile'))

    docx_stream = export_to_docx(reconciliation_explanation_markdown, RECONCILIATION_DOCX_FILENAME)
    if docx_stream:
        return send_file(
            docx_stream,
            as_attachment=True,
            download_name=RECONCILIATION_DOCX_FILENAME,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    else:
        flash('Error exporting reconciliation explanation to DOCX.', 'error')
        return redirect(url_for('reconcile'))


@app.route('/admin')
@login_required
@admin_required
def admin():
    """Admin dashboard for managing users and system settings."""
    api_status, api_message = test_api_connection()
    
    # Get system information
    system_info = {
        'api_key_configured': bool(API_KEY),
        'api_key_status': "Valid" if api_status else "Invalid or connection error",
        'api_message': api_message,
        'upload_folder': UPLOAD_FOLDER,
        'upload_folder_exists': os.path.exists(UPLOAD_FOLDER),
        'upload_folder_writable': os.access(UPLOAD_FOLDER, os.W_OK) if os.path.exists(UPLOAD_FOLDER) else False,
        'allowed_extensions': ", ".join(ALLOWED_EXTENSIONS),
        'models_available': [DEFAULT_MODEL_NAME, THINKING_MODEL_NAME],
        'user_count': len(users)
    }
    
    # Get user list for admin
    user_list = []
    for user_id, user_data in users.items():
        user_list.append({
            'id': user_id,
            'username': user_data['username'],
            'is_admin': user_data.get('is_admin', False)
        })
    
    return render_template('admin.html', system_info=system_info, users=user_list, current_user=current_user)


@app.route('/admin/test_api', methods=['POST'])
@login_required
@admin_required
def test_api():
    """API endpoint to test the API connection."""
    status, message = test_api_connection()
    return jsonify({
        'status': status,
        'message': message
    })


if __name__ == '__main__':
    if configure_api():
        app.run(debug=True, host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))
